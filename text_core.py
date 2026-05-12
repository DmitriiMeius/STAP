from __future__ import annotations

import csv
import re
import zipfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import fitz
import pandas as pd
from bs4 import BeautifulSoup
from charset_normalizer import from_bytes
from docx import Document
from ebooklib import epub
from striprtf.striprtf import rtf_to_text


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".rtf",
    ".pdf",
    ".docx",
    ".xls",
    ".xlsx",
    ".fb2",
    ".zip",
    ".epub",
    ".html",
    ".htm",
    ".md",
    ".csv",
}

WORD_RE = re.compile(r"[A-Za-zА-Яа-яЁё0-9]+(?:[-'][A-Za-zА-Яа-яЁё0-9]+)?")
SENTENCE_RE = re.compile(r"[^.!?。！？]+[.!?。！？]*", re.MULTILINE)

DEFAULT_STOP_WORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "has",
    "he",
    "in",
    "is",
    "it",
    "its",
    "of",
    "on",
    "that",
    "the",
    "to",
    "was",
    "were",
    "will",
    "with",
    "в",
    "во",
    "и",
    "к",
    "ко",
    "на",
    "не",
    "но",
    "о",
    "об",
    "от",
    "по",
    "с",
    "со",
    "у",
    "за",
    "из",
    "для",
    "как",
    "что",
    "это",
    "или",
    "а",
    "же",
    "ли",
    "бы",
}


@dataclass(frozen=True)
class TextMetrics:
    characters: int
    characters_no_spaces: int
    words: int
    unique_words: int
    sentences: int
    paragraphs: int
    average_word_length: float
    average_sentence_words: float
    top_words: list[tuple[str, int]]


@dataclass(frozen=True)
class FileAnalysis:
    path: Path
    text: str
    metrics: TextMetrics


def is_supported(path: Path) -> bool:
    if path.suffix.lower() == ".zip" and path.name.lower().endswith(".fb2.zip"):
        return True
    return path.suffix.lower() in SUPPORTED_EXTENSIONS


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    detected = from_bytes(raw).best()
    if detected is None:
        return raw.decode("utf-8", errors="replace")
    return str(detected)


def extract_pdf(path: Path) -> str:
    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_spreadsheet(path: Path) -> str:
    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    parts: list[str] = []
    for sheet_name, frame in sheets.items():
        parts.append(f"[{sheet_name}]")
        frame = frame.fillna("")
        parts.extend(
            " | ".join(str(value).strip() for value in row if str(value).strip())
            for row in frame.to_numpy()
        )
    return "\n".join(parts)


def extract_csv(path: Path) -> str:
    text = read_text_file(path)
    rows = csv.reader(text.splitlines())
    return "\n".join(" | ".join(cell.strip() for cell in row if cell.strip()) for row in rows)


def extract_html(path: Path) -> str:
    soup = BeautifulSoup(read_text_file(path), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n")


def extract_fb2(path: Path) -> str:
    soup = BeautifulSoup(read_text_file(path), "xml")
    return soup.get_text("\n")


def extract_fb2_zip(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        fb2_names = [name for name in archive.namelist() if name.lower().endswith(".fb2")]
        if not fb2_names:
            return ""
        raw = archive.read(fb2_names[0])
    detected = from_bytes(raw).best()
    text = str(detected) if detected is not None else raw.decode("utf-8", errors="replace")
    soup = BeautifulSoup(text, "xml")
    return soup.get_text("\n")


def extract_epub(path: Path) -> str:
    book = epub.read_epub(str(path))
    parts: list[str] = []
    for item in book.get_items():
        if item.get_type() == 9:
            soup = BeautifulSoup(item.get_content(), "html.parser")
            parts.append(soup.get_text("\n"))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if path.name.lower().endswith(".fb2.zip"):
        return clean_text(extract_fb2_zip(path))
    if suffix in {".txt", ".md"}:
        return clean_text(read_text_file(path))
    if suffix == ".rtf":
        return clean_text(rtf_to_text(read_text_file(path)))
    if suffix == ".pdf":
        return clean_text(extract_pdf(path))
    if suffix == ".docx":
        return clean_text(extract_docx(path))
    if suffix in {".xls", ".xlsx"}:
        return clean_text(extract_spreadsheet(path))
    if suffix == ".csv":
        return clean_text(extract_csv(path))
    if suffix in {".html", ".htm"}:
        return clean_text(extract_html(path))
    if suffix == ".fb2":
        return clean_text(extract_fb2(path))
    if suffix == ".epub":
        return clean_text(extract_epub(path))
    raise ValueError(f"Unsupported file format: {path.suffix}")


def tokenize_words(text: str) -> list[str]:
    return [match.group(0).lower() for match in WORD_RE.finditer(text)]


def count_paragraphs(text: str) -> int:
    return len([part for part in re.split(r"\n\s*\n", text) if part.strip()])


def count_sentences(text: str) -> int:
    return len([part for part in SENTENCE_RE.findall(text) if part.strip()])


def analyze_text(text: str, *, top_limit: int = 25, stop_words: Iterable[str] = DEFAULT_STOP_WORDS) -> TextMetrics:
    words = tokenize_words(text)
    filtered_words = [word for word in words if word not in stop_words and len(word) > 1]
    word_counter = Counter(filtered_words)
    sentence_count = count_sentences(text)
    word_count = len(words)
    total_word_chars = sum(len(word) for word in words)

    return TextMetrics(
        characters=len(text),
        characters_no_spaces=len(re.sub(r"\s+", "", text)),
        words=word_count,
        unique_words=len(set(words)),
        sentences=sentence_count,
        paragraphs=count_paragraphs(text),
        average_word_length=(total_word_chars / word_count) if word_count else 0.0,
        average_sentence_words=(word_count / sentence_count) if sentence_count else 0.0,
        top_words=word_counter.most_common(top_limit),
    )


def analyze_file(path: Path) -> FileAnalysis:
    text = extract_text(path)
    return FileAnalysis(path=path, text=text, metrics=analyze_text(text))


def combine_analyses(analyses: Iterable[FileAnalysis]) -> TextMetrics:
    return analyze_text("\n\n".join(item.text for item in analyses))


def format_metrics(metrics: TextMetrics) -> str:
    lines = [
        f"Characters: {metrics.characters}",
        f"Characters without spaces: {metrics.characters_no_spaces}",
        f"Words: {metrics.words}",
        f"Unique words: {metrics.unique_words}",
        f"Sentences: {metrics.sentences}",
        f"Paragraphs: {metrics.paragraphs}",
        f"Average word length: {metrics.average_word_length:.2f}",
        f"Average sentence length: {metrics.average_sentence_words:.2f} words",
        "",
        "Top words:",
    ]
    if metrics.top_words:
        lines.extend(f"{word}: {count}" for word, count in metrics.top_words)
    else:
        lines.append("No words found.")
    return "\n".join(lines)
